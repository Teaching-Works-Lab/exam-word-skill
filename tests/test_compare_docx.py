from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document
from docx.shared import Cm, Pt

from conftest import load_script


compare_docx = load_script("compare_docx")


def make_docx(path: Path, text: str, *, left_margin_cm: float = 3.0, font_size: int = 12):
    document = Document()
    document.sections[0].left_margin = Cm(left_margin_cm)
    document.styles["Normal"].font.size = Pt(font_size)
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "题号"
    table.cell(0, 1).text = "分数"
    document.save(path)


class CompareDocxTests(unittest.TestCase):
    def test_separates_content_and_formatting_differences(self):
        with TemporaryDirectory() as directory:
            temp_path = Path(directory)
            left = temp_path / "标准.docx"
            right = temp_path / "试卷.docx"
            make_docx(left, "相同题目", left_margin_cm=3.0, font_size=12)
            make_docx(right, "相同题目", left_margin_cm=2.0, font_size=11)

            report = compare_docx.compare_documents(left, right)

            self.assertFalse(report["content"]["different"])
            self.assertTrue(report["formatting"]["different"])
            self.assertTrue(report["formatting"]["page_setup"])
            self.assertTrue(report["formatting"]["styles"])
            self.assertEqual(report["summary"]["content_difference_count"], 0)


    def test_reports_text_difference_without_format_drift(self):
        with TemporaryDirectory() as directory:
            temp_path = Path(directory)
            left = temp_path / "标准.docx"
            right = temp_path / "试卷.docx"
            make_docx(left, "第一道题")
            make_docx(right, "第二道题")

            report = compare_docx.compare_documents(left, right)

            self.assertTrue(report["content"]["different"])
            self.assertFalse(report["formatting"]["different"])
            self.assertTrue(report["content"]["paragraphs"])
            self.assertIn("内容差异", compare_docx.render_markdown(report))


if __name__ == "__main__":
    unittest.main()
