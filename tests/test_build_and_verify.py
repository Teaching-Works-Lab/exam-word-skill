import hashlib
import json
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from zipfile import ZipFile

from docx import Document

from conftest import ROOT, load_script


build_from_spec = load_script("build_from_spec")
verify_docx = load_script("verify_docx")


PRESERVED_PARTS = {
    "word/styles.xml",
    "word/numbering.xml",
    "word/theme/theme1.xml",
    "word/footer1.xml",
    "word/fontTable.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
}


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def part_digests(path: Path):
    with ZipFile(path) as archive:
        return {
            part: hashlib.sha256(archive.read(part)).hexdigest()
            for part in PRESERVED_PARTS
            if part in archive.namelist()
        }


class BuildAndVerifyTests(unittest.TestCase):
    def test_builds_from_reference_without_mutating_template(self):
        with TemporaryDirectory() as directory:
            temp_path = Path(directory)
            template = ROOT / "assets" / "reference.docx"
            before = digest(template)
            spec_path = temp_path / "试卷结构.json"
            output = temp_path / "标准试卷.docx"
            spec_path.write_text(
                json.dumps(
                    {
                        "metadata": {
                            "school": "沧州交通学院",
                            "academic_term": "2025－2026学年第二学期",
                            "unit": "机械学院",
                            "teacher": "测试教师",
                            "course": "测试课程",
                            "paper": "A 卷",
                            "exam_mode": "闭卷",
                            "summary": "本试卷共有一道大题（满分：100 分）",
                        },
                        "pages": [
                            {
                                "blocks": [
                                    {"type": "heading", "text": "一、测试题（满分100分，共1题）"},
                                    {"type": "question", "text": "1、这是测试题目。（100分）"},
                                ]
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = build_from_spec.build_document(template, spec_path, output)

            self.assertTrue(output.exists())
            self.assertEqual(digest(template), before)
            self.assertEqual(part_digests(output), part_digests(template))
            self.assertEqual(result["output"], str(output))
            self.assertIn("这是测试题目", "\n".join(p.text for p in Document(output).paragraphs))

            verification = verify_docx.verify_document(output, template)
            self.assertTrue(verification["valid"])
            self.assertTrue(verification["preserved_parts_match"])


    def test_verifier_rejects_missing_docx(self):
        with TemporaryDirectory() as directory:
            template = ROOT / "assets" / "reference.docx"
            result = verify_docx.verify_document(Path(directory) / "不存在.docx", template)
            self.assertFalse(result["valid"])
            self.assertIn("output_missing", result["errors"])

    def test_score_table_clears_unspecified_template_cells(self):
        with TemporaryDirectory() as directory:
            temp_path = Path(directory)
            template = ROOT / "assets" / "reference.docx"
            spec_path = temp_path / "spec.json"
            output = temp_path / "output.docx"
            spec_path.write_text(
                json.dumps(
                    {
                        "metadata": {
                            "school": "沧州交通学院",
                            "academic_term": "2025－2026学年第二学期",
                            "unit": "机械学院",
                            "teacher": "测试教师",
                            "course": "测试课程",
                            "summary": "本试卷共有两道大题（满分：100 分）",
                        },
                        "score_table": [
                            ["题号", "一", "二", "总分"],
                            ["分数", "40", "60", "100"],
                            ["得分", "", "", ""],
                        ],
                        "pages": [{"blocks": [{"type": "text", "text": "测试"}]}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            build_from_spec.build_document(template, spec_path, output)
            table = Document(output).tables[0]

            self.assertEqual([cell.text for cell in table.rows[0].cells][4:], ["", "", "", ""])


if __name__ == "__main__":
    unittest.main()
