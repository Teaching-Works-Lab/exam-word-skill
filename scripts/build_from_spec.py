import argparse
import json
from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips


PRESERVED_PARTS = {
    "word/styles.xml",
    "word/numbering.xml",
    "word/theme/theme1.xml",
    "word/footer1.xml",
    "word/fontTable.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
}
BLOCK_TYPES = {"heading", "question", "option", "option-row", "text", "spacer", "table"}


def _clear_body(document):
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _configure_paragraph(paragraph, block):
    alignments = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    if block.get("align"):
        paragraph.alignment = alignments[block["align"]]
    formatting = paragraph.paragraph_format
    formatting.space_before = Pt(block.get("before_pt", 0))
    formatting.space_after = Pt(block.get("after_pt", 0))
    formatting.line_spacing = Pt(block.get("line_pt", 20))
    formatting.keep_with_next = bool(block.get("keep_next", False))
    if "left_twips" in block:
        formatting.left_indent = Twips(block["left_twips"])
    if "first_twips" in block:
        formatting.first_line_indent = Twips(block["first_twips"])


def _add_text(document, block):
    style = block.get("style", "Normal")
    paragraph = document.add_paragraph(style=style)
    _configure_paragraph(paragraph, block)
    run = paragraph.add_run(block.get("text", ""))
    run.bold = block.get("bold")
    return paragraph


def _add_tabbed_text(document, segments, *, left_twips=425):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Twips(left_twips)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(20)
    if len(segments) >= 4:
        stops = (1.62, 3.15, 4.72)
    elif len(segments) == 2:
        stops = (3.1,)
    else:
        stops = ()
    for position in stops:
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(position), WD_TAB_ALIGNMENT.LEFT)
    for index, segment in enumerate(segments):
        if index:
            paragraph.add_run("\t")
        paragraph.add_run(str(segment))


def _add_header(document, metadata):
    title = document.add_paragraph(style="Heading 1")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.add_run(metadata["school"])

    term = document.add_paragraph(style="Heading 2")
    term.paragraph_format.line_spacing = Pt(20)
    term.paragraph_format.tab_stops.add_tab_stop(Inches(3.4), WD_TAB_ALIGNMENT.CENTER)
    term.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
    term.add_run(metadata["academic_term"])
    term.add_run("\t")
    term.add_run(f"命题单位：{metadata['unit']}")
    term.add_run("\t")
    term.add_run(f"命题教师：{metadata['teacher']}")

    course = document.add_paragraph(style="Normal")
    course.paragraph_format.line_spacing = Pt(20)
    course.paragraph_format.tab_stops.add_tab_stop(Inches(4.15), WD_TAB_ALIGNMENT.CENTER)
    course.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
    first = course.add_run(f"【{metadata['course']}】课程")
    first.bold = True
    course.add_run("\t")
    course.add_run(f"（{metadata.get('paper', 'A 卷')}）")
    course.add_run("\t")
    last = course.add_run(f"【{metadata.get('exam_mode', '闭卷')}】")
    last.bold = True

    information = document.add_paragraph(style="Heading 2")
    information.paragraph_format.line_spacing = Pt(20)
    for index, label in enumerate(("序号", "学号", "姓名", "专业及班级")):
        if index:
            information.add_run("  ")
        information.add_run(label)
        underline = information.add_run(" " * 13)
        underline.underline = True
    _add_text(document, {"text": metadata["summary"], "style": "Heading 2"})


def _populate_score_table(score_table, score_rows):
    if not score_rows:
        return
    for row_index, row in enumerate(score_table.rows):
        row_values = score_rows[row_index] if row_index < len(score_rows) else []
        for column_index, cell in enumerate(row.cells):
            value = row_values[column_index] if column_index < len(row_values) else ""
            cell.text = str(value)


def _add_block(document, block):
    block_type = block["type"]
    if block_type == "heading":
        data = {"bold": True, "keep_next": True, "before_pt": 4, **block}
        _add_text(document, data)
    elif block_type == "question":
        data = {"left_twips": 283, "first_twips": -283, **block}
        _add_text(document, data)
    elif block_type == "option":
        data = {"left_twips": 425, **block}
        _add_text(document, data)
    elif block_type == "option-row":
        _add_tabbed_text(document, block["items"], left_twips=block.get("left_twips", 425))
    elif block_type in {"text", "spacer"}:
        _add_text(document, block)
    elif block_type == "table":
        rows = block.get("rows", [])
        if not rows:
            return
        table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
        if block.get("style"):
            table.style = block["style"]
        for row_index, values in enumerate(rows):
            for column_index, value in enumerate(values):
                table.cell(row_index, column_index).text = str(value)


def _validate_spec(spec):
    metadata = spec.get("metadata", {})
    required = {"school", "academic_term", "unit", "teacher", "course", "summary"}
    missing = sorted(required - set(metadata))
    if missing:
        raise ValueError(f"Missing metadata fields: {', '.join(missing)}")
    pages = spec.get("pages")
    if not isinstance(pages, list) or not pages:
        raise ValueError("At least one page is required")
    for page in pages:
        for block in page.get("blocks", []):
            if block.get("type") not in BLOCK_TYPES:
                raise ValueError(f"Unsupported block type: {block.get('type')}")


def _restore_parts(template_path: Path, output_path: Path):
    temporary_path = output_path.with_suffix(".tmp.docx")
    with ZipFile(template_path) as template_archive, ZipFile(output_path) as output_archive:
        preserved = {
            part: template_archive.read(part)
            for part in PRESERVED_PARTS
            if part in template_archive.namelist()
        }
        with ZipFile(temporary_path, "w") as temporary_archive:
            written = set()
            for info in output_archive.infolist():
                data = preserved.get(info.filename, output_archive.read(info.filename))
                temporary_archive.writestr(info, data)
                written.add(info.filename)
            for part, data in preserved.items():
                if part not in written:
                    temporary_archive.writestr(part, data)
    temporary_path.replace(output_path)
    return sorted(preserved)


def build_document(template_path, spec_path, output_path):
    template_path = Path(template_path)
    spec_path = Path(spec_path)
    output_path = Path(output_path)
    if not template_path.exists():
        raise FileNotFoundError(template_path)
    spec = json.loads(spec_path.read_text(encoding="utf-8-sig"))
    _validate_spec(spec)

    document = Document(template_path)
    score_table_element = deepcopy(document.tables[0]._tbl) if document.tables else None
    _clear_body(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    update_fields = document.settings._element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        document.settings._element.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    _add_header(document, spec["metadata"])
    if score_table_element is not None:
        document._body._element.insert(len(document._body._element) - 1, score_table_element)
        _populate_score_table(document.tables[0], spec.get("score_table"))

    for page_index, page in enumerate(spec["pages"]):
        if page_index:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run().add_break(WD_BREAK.PAGE)
        for block in page.get("blocks", []):
            _add_block(document, block)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    preserved = _restore_parts(template_path, output_path)
    return {"output": str(output_path), "preserved_parts": preserved, "page_blocks": len(spec["pages"])}


def main() -> None:
    parser = argparse.ArgumentParser(description="Build an editable exam DOCX from structured JSON")
    parser.add_argument("template")
    parser.add_argument("spec")
    parser.add_argument("output")
    args = parser.parse_args()
    print(json.dumps(build_document(args.template, args.spec, args.output), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
