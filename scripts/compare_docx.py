import argparse
import difflib
import hashlib
import json
from pathlib import Path
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

from docx import Document


WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TEXT_TAGS = {f"{{{WORD_NAMESPACE}}}t", f"{{{WORD_NAMESPACE}}}instrText"}


def _length(value):
    return int(value) if value is not None else None


def _enum(value):
    return int(value) if value is not None else None


def _hash_part(path: Path, part: str, *, strip_text: bool = False):
    with ZipFile(path) as archive:
        if part not in archive.namelist():
            return None
        data = archive.read(part)
    if strip_text:
        root = ElementTree.fromstring(data)
        for element in root.iter():
            if element.tag in TEXT_TAGS:
                element.text = ""
        data = ElementTree.tostring(root, encoding="utf-8")
    return hashlib.sha256(data).hexdigest()


def _paragraph_content(document: Document):
    return [paragraph.text for paragraph in document.paragraphs]


def _table_content(document: Document):
    return [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]


def _paragraph_format(document: Document):
    records = []
    for paragraph in document.paragraphs:
        formatting = paragraph.paragraph_format
        records.append(
            {
                "style": paragraph.style.style_id if paragraph.style else None,
                "alignment": _enum(paragraph.alignment),
                "left_indent": _length(formatting.left_indent),
                "right_indent": _length(formatting.right_indent),
                "first_line_indent": _length(formatting.first_line_indent),
                "space_before": _length(formatting.space_before),
                "space_after": _length(formatting.space_after),
                "line_spacing": _length(formatting.line_spacing) if hasattr(formatting.line_spacing, "__int__") else formatting.line_spacing,
                "keep_with_next": formatting.keep_with_next,
                "keep_together": formatting.keep_together,
                "page_break_before": formatting.page_break_before,
                "runs": [
                    {
                        "style": run.style.style_id if run.style else None,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": str(run.underline) if run.underline is not None else None,
                        "font_name": run.font.name,
                        "font_size": _length(run.font.size),
                    }
                    for run in paragraph.runs
                ],
            }
        )
    return records


def _table_format(document: Document):
    return [
        {
            "style": table.style.style_id if table.style else None,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "grid_widths": [_length(column.width) for column in table.columns],
        }
        for table in document.tables
    ]


def _page_setup(document: Document):
    attributes = (
        "page_width",
        "page_height",
        "left_margin",
        "right_margin",
        "top_margin",
        "bottom_margin",
        "header_distance",
        "footer_distance",
        "gutter",
    )
    return [
        {attribute: _length(getattr(section, attribute)) for attribute in attributes}
        | {"orientation": _enum(section.orientation), "start_type": _enum(section.start_type)}
        for section in document.sections
    ]


def _differences(left, right):
    maximum = max(len(left), len(right))
    return [
        {
            "index": index,
            "left": left[index] if index < len(left) else None,
            "right": right[index] if index < len(right) else None,
        }
        for index in range(maximum)
        if (left[index] if index < len(left) else None) != (right[index] if index < len(right) else None)
    ]


def compare_documents(left_path, right_path):
    left_path = Path(left_path)
    right_path = Path(right_path)
    for path in (left_path, right_path):
        if not path.exists():
            raise FileNotFoundError(path)
        try:
            with ZipFile(path) as archive:
                archive.testzip()
        except BadZipFile as error:
            raise ValueError(f"Invalid DOCX package: {path}") from error

    left = Document(left_path)
    right = Document(right_path)
    paragraph_differences = _differences(_paragraph_content(left), _paragraph_content(right))
    table_content_differences = _differences(_table_content(left), _table_content(right))
    page_setup_differences = _differences(_page_setup(left), _page_setup(right))
    paragraph_format_differences = _differences(_paragraph_format(left), _paragraph_format(right))
    table_format_differences = _differences(_table_format(left), _table_format(right))

    package_checks = {}
    for label, part, strip_text in (
        ("styles", "word/styles.xml", False),
        ("numbering", "word/numbering.xml", False),
        ("headers", "word/header1.xml", True),
        ("footers", "word/footer1.xml", True),
    ):
        left_hash = _hash_part(left_path, part, strip_text=strip_text)
        right_hash = _hash_part(right_path, part, strip_text=strip_text)
        if left_hash != right_hash:
            package_checks[label] = {"left": left_hash, "right": right_hash}

    content_count = len(paragraph_differences) + len(table_content_differences)
    formatting_count = (
        len(page_setup_differences)
        + len(paragraph_format_differences)
        + len(table_format_differences)
        + len(package_checks)
    )
    return {
        "left": str(left_path),
        "right": str(right_path),
        "content": {
            "different": content_count > 0,
            "paragraphs": paragraph_differences,
            "tables": table_content_differences,
        },
        "formatting": {
            "different": formatting_count > 0,
            "page_setup": page_setup_differences,
            "paragraphs": paragraph_format_differences,
            "tables": table_format_differences,
            **package_checks,
        },
        "summary": {
            "content_difference_count": content_count,
            "formatting_difference_count": formatting_count,
        },
    }


def render_markdown(report) -> str:
    summary = report["summary"]
    lines = [
        "# Word 对比报告",
        "",
        f"- 左侧文件：`{report['left']}`",
        f"- 右侧文件：`{report['right']}`",
        f"- 内容差异：{summary['content_difference_count']} 项",
        f"- 格式差异：{summary['formatting_difference_count']} 项",
        "",
        "## 内容差异",
        "",
    ]
    if report["content"]["different"]:
        for item in report["content"]["paragraphs"]:
            lines.append(f"- 段落 {item['index'] + 1}：`{item['left']}` → `{item['right']}`")
        for item in report["content"]["tables"]:
            lines.append(f"- 表格 {item['index'] + 1} 的单元格内容不同。")
    else:
        lines.append("- 未发现内容差异。")
    lines.extend(["", "## 格式差异", ""])
    if report["formatting"]["different"]:
        labels = {
            "page_setup": "页面设置",
            "paragraphs": "段落或字符格式",
            "tables": "表格结构或尺寸",
            "styles": "样式定义",
            "numbering": "编号定义",
            "headers": "页眉结构",
            "footers": "页脚结构",
        }
        for key, label in labels.items():
            if report["formatting"].get(key):
                lines.append(f"- {label}不同。")
    else:
        lines.append("- 未发现格式差异。")
    return "\n".join(lines) + "\n"


def main() -> None:
    parser = argparse.ArgumentParser(description="Compare two DOCX files")
    parser.add_argument("left")
    parser.add_argument("right")
    parser.add_argument("--json-output")
    parser.add_argument("--markdown-output")
    args = parser.parse_args()
    report = compare_documents(args.left, args.right)
    if args.json_output:
        Path(args.json_output).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if args.markdown_output:
        Path(args.markdown_output).write_text(render_markdown(report), encoding="utf-8")
    if not args.json_output and not args.markdown_output:
        print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
